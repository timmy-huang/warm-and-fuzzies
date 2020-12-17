from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, RGBColor, Inches
import random
import json

# Thanks to the emoji library from @chazgiese found in emojigen
# https://github.com/chazgiese/EmojiGen/blob/master/emojigen.js
# as a part of the EmojiGen project, https://github.com/chazgiese/EmojiGen

EMOJIS = [
  "😀", "😬", "😁", "😂", "😃",
  "😄", "😅", "😆", "😇", "😉",
  "😊", "🙂", "🙃", "☺️", "😋",
  "😌", "😍", "😘", "😗", "😙",
  "😚", "😜", "😝", "😛", "🤑",
  "🤓", "😎", "🤗", "😏", "😶",
  "😐", "😑", "😒", "🙄", "🤔",
  "😳", "😞", "😟", "😠", "😡",
  "😔", "😕", "🙁", "☹️", "😣",
  "😖", "😫", "😩", "😤", "😮",
  "😱", "😨", "😰", "😯", "😦",
  "😧", "😢", "😥", "😪", "😓",
  "😭", "😵", "😲", "🤐", "😷",
  "🤒", "🤕", "😴", "💤", "💩",
  "😈", "👿", "👹", "👺", "💀",
  "👻", "👽", "🤖", "😺", "😸",
  "😹", "😻", "😼", "😽", "🙀",
  "😿", "😾", "🙌", "👏", "👋",
  "👍", "👎", "👊", "✊", "✌️",
  "👌", "✋", "👐", "💪", "🙏",
  "☝️", "👆", "👇", "👈", "👉",
  "🖕", "🖐", "🤘", "🖖", "✍️",
  "💅", "👄", "👅", "👂", "👃",
  "👁", "👀", "👤", "👥", "🗣",
  "👶", "👦", "👧", "👨", "👩",
  "👱", "👴", "👵", "👲", "👳",
  "👮", "👷", "💂", "🕵", "🎅",
  "👼", "👸", "👰", "🚶", "🏃",
  "💃", "👯", "👫", "👬", "👭",
  "🙇", "💁", "🙅", "🙆", "🙋",
  "🙎", "🙍", "💇", "💆", "💑",

  "👩‍❤️‍👩",
  "👨‍❤️‍👨", "💏",
  "👩‍❤️‍💋‍👩",
  "👨‍❤️‍💋‍👨",

  "👪", "👨‍👩‍👧",
  "👨‍👩‍👧‍👦",
  "👨‍👩‍👦‍👦",
  "👨‍👩‍👦‍👦",
  "👨‍👩‍👧‍👧",
  "👩‍👩‍👦",
  "👩‍👩‍👧",
  "👩‍👩‍👧‍👦",
  "👩‍👩‍👦‍👦",
  "👩‍👩‍👧‍👧",
  "👨‍👨‍👦",
  "👨‍👨‍👧",
  "👨‍👨‍👧‍👦",
  "👨‍👨‍👦‍👦",
  "👨‍👨‍👧‍👧",

  "👚", "👕", "👖", "👔", "👗",
  "👙", "👘", "💄", "💋", "👣",
  "👠", "👡", "👢", "👞", "👟",
  "👒", "🎩", "⛑", "🎓", "👑",
  "🎒", "👝", "👛", "👜", "💼",
  "👓", "🕶", "💍", "🌂",


  "🐶", "🐱", "🐭", "🐹", "🐰",
  "🐻", "🐼", "🐨", "🐯", "🦁",
  "🐮", "🐷", "🐽", "🐸", "🐙",
  "🐵", "🙈", "🙉", "🙊", "🐒",
  "🐔", "🐧", "🐦", "🐤", "🐣",
  "🐥", "🐺", "🐗", "🐴", "🦄",
  "🐝", "🐛", "🐌", "🐞", "🐜",
  "🕷", "🦂", "🦀", "🐍", "🐢",
  "🐠", "🐟", "🐡", "🐬", "🐳",
  "🐋", "🐊", "🐆", "🐅", "🐃",
  "🐂", "🐄", "🐪", "🐫", "🐘",
  "🐐", "🐏", "🐑", "🐎", "🐖",
  "🐀", "🐁", "🐓", "🦃", "🕊",
  "🐕", "🐩", "🐈", "🐇", "🐿",
  "🐾", "🐉", "🐲", "🌵", "🎄",
  "🌲", "🌳", "🌴", "🌱", "🌿",
  "☘", "🍀", "🎍", "🎋", "🍃",
  "🍂", "🍁", "🌾", "🌺", "🌻",
  "🌹", "🌷", "🌼", "🌸", "💐",
  "🍄", "🌰", "🎃", "🐚", "🕸",
  "🌎", "🌍", "🌏", "🌕", "🌖",
  "🌗", "🌘", "🌑", "🌒", "🌓",
  "🌔", "🌚", "🌝", "🌛", "🌜",
  "🌞", "🌙", "⭐️", "🌟", "💫",
  "✨", "☄️", "☀️", "🌤",
  "⛅️",
  "🌥", "🌦", "☁️", "🌧", "⛈",
  "🌩", "⚡️", "🔥", "💥",
  "❄️",
  "🌨", "☃️", "⛄️", "🌬",
  "💨",
  "🌪", "🌫", "☂️", "☔️",
  "💧",
  "💦", "🌊",

  "🍏", "🍎", "🍐", "🍊", "🍋",
  "🍌", "🍉", "🍇", "🍓", "🍈",
  "🍒", "🍑", "🍍", "🍅", "🍆",
  "🌶", "🌽", "🍠", "🍯", "🍞",
  "🧀", "🍗", "🍖", "🍤", "🍳",
  "🍔", "🍟", "🌭", "🍕", "🍝",
  "🌮", "🌯", "🍜", "🍲", "🍥",
  "🍣", "🍱", "🍛", "🍙", "🍚",
  "🍘", "🍢", "🍡", "🍧", "🍨",
  "🍦", "🍰", "🎂", "🍮", "🍬",
  "🍭", "🍫", "🍿", "🍩", "🍪",
  "🍺", "🍻", "🍷", "🍸", "🍹",
  "🍾", "🍶", "🍵", "☕️", "🍼",
  "🍴", "🍽",

  "⚽️", "🏀", "🏈", "⚾️",
  "🎾",
  "🏐", "🏉", "🎱", "⛳️", "🏌",
  "🏓", "🏸", "🏒", "🏑", "🏏",
  "🎿", "⛷", "🏂", "⛸", "🏹",
  "🎣", "🚣", "🏊", "🏄", "🛀",
  "⛹", "🏋", "🚴", "🚵", "🏇",
  "🕴", "🏆", "🎽", "🏅", "🎖",
  "🎗", "🏵", "🎫", "🎟", "🎭",
  "🎨", "🎪", "🎤", "🎧", "🎼",
  "🎹", "🎷", "🎺", "🎸", "🎻",
  "🎬", "🎮", "👾", "🎯", "🎲",
  "🎰", "🎳",

  "🚗", "🚕", "🚙", "🚌", "🚎",
  "🏎", "🚓", "🚑", "🚒", "🚐",
  "🚚", "🚛", "🚜", "🏍", "🚲",
  "🚨", "🚔", "🚍", "🚘", "🚖",
  "🚡", "🚠", "🚟", "🚃", "🚋",
  "🚝", "🚄", "🚅", "🚈", "🚞",
  "🚂", "🚆", "🚇", "🚊", "🚉",
  "🚁", "🛩", "✈️", "🛫", "🛬",
  "⛵️", "🛥", "🚤", "⛴", "🛳",
  "🚀", "🛰", "💺", "⚓️", "🚧",
  "⛽️", "🚏", "🚦", "🚥", "🏁",
  "🚢", "🎡", "🎢", "🎠", "🏗",
  "🌁", "🗼", "🏭", "⛲️", "🎑",
  "⛰", "🏔", "🗻", "🌋", "🗾",
  "🏕", "⛺️", "🏞", "🛣", "🛤",
  "🌅", "🌄", "🏜", "🏖", "🏝",
  "🌇", "🌆", "🏙", "🌃", "🌉",
  "🌌", "🌠", "🎇", "🎆", "🌈",
  "🏘", "🏰", "🏯", "🏟", "🗽",
  "🏠", "🏡", "🏚", "🏢", "🏬",
  "🏣", "🏤", "🏥", "🏦", "🏨",
  "🏪", "🏫", "🏩", "💒", "🏛",
  "⛪️", "🕌", "🕍", "🕋", "⛩",

  "⌚️", "📱", "📲", "💻", "⌨️",
  "🖥", "🖨", "🖱", "🖲", "🕹",
  "🗜", "💽", "💾", "💿", "📀",
  "📼", "📷", "📸", "📹", "🎥",
  "📽", "🎞", "📞", "☎️", "📟",
  "📠", "📺", "📻", "🎙", "🎚",
  "🎛", "⏱", "⏲", "⏰", "🕰",
  "⏳", "⌛️", "📡", "🔋", "🔌",
  "💡", "🔦", "🕯", "🗑", "🛢",
  "💸", "💵", "💴", "💶", "💷",
  "💰", "💳", "💎", "⚖", "🔧",
  "🔨", "⚒", "🛠", "⛏", "🔩",
  "⚙", "⛓", "💣", "🔪",
  "🗡", "⚔", "🛡", "🚬", "☠️",
  "⚰", "⚱", "🏺", "🔮", "📿",
  "💈", "⚗", "🔭", "🔬", "🕳",
  "💊", "💉", "🌡", "🏷", "🔖",
  "🚽", "🚿", "🛁", "🔑", "🗝",
  "🛋", "🛌", "🛏", "🚪", "🛎",
  "🖼", "🗺", "⛱", "🗿", "🛍",
  "🎈", "🎏", "🎀", "🎁", "🎊",
  "🎉", "🎎", "🎐", "🎌", "🏮",
  "✉️", "📩", "📨", "📧", "💌",
  "📮", "📪", "📫", "📬", "📭",
  "📦", "📯", "📥", "📤", "📜",
  "📃", "📑", "📊", "📈", "📉",
  "📄", "📅", "📆", "🗓", "📇",
  "🗃", "🗳", "🗄", "📋", "🗒",
  "📁", "📂", "🗂", "🗞", "📰",
  "📓", "📕", "📗", "📘", "📙",
  "📔", "📒", "📚", "📖", "🔗",
  "📎", "🖇", "✂️", "📐", "📏",
  "📌", "📍", "🚩", "🏳", "🏴",
  "🔐", "🔒", "🔓", "🔏", "🖊",
  "🖋", "✒️", "📝", "✏️", "🖍",
  "🖌", "🔍", "🔎",

  "❤️", "💛", "💚", "💙", "💜",
  "💔", "❣️", "💕", "💞", "💓",
  "💗", "💖", "💘", "💝", "💟",
  "☮️", "✝️", "☪️", "🕉",
  "☸️",
  "✡️", "🔯", "🕎", "☯️",
  "☦️",
  "🛐", "⛎", "♈️", "♉️",
  "♊️",
  "♋️", "♌️", "♍️", "♎️",
  "♏️",
  "♐️", "♑️", "♒️", "♓️",
  "🆔",
  "⚛",

  "🈳",

  "🈹", "☢️", "☣️",
  "📴", "📳",

  "🈶", "🈚️", "🈸", "🈺", "🈷️",

  "✴️", "🆚",
  "🉑",
  "💮",

  "🉐", "㊙️", "㊗️", "🈴", "🈵", "🈲",

  "🅰️", "🅱️",
  "🆎",
  "🆑", "🅾️", "🆘", "⛔️",
  "📛",
  "🚫", "❌", "⭕️", "💢", "♨️",
  "🚷", "🚯", "🚳", "🚱", "🔞",
  "📵", "❗️", "❕", "❓", "❔",
  "‼️", "⁉️", "💯", "🔅", "🔆",
  "🔱", "⚜", "〽️", "⚠️", "🚸",
  "🔰", "♻️",


  "🈯️",

  "💹",
  "❇️",
  "✳️", "❎", "✅", "💠", "🌀",
  "➿", "🌐", "Ⓜ️", "🏧", "🈂️",
  "🛂", "🛃", "🛄", "🛅", "♿️",
  "🚭", "🚾", "🅿️", "🚰", "🚹",
  "🚺", "🚼", "🚻", "🚮", "🎦",
  "📶", "🈁", "🆖", "🆗", "🆙",
  "🆒", "🆕", "🆓", "0️⃣",
  "1️⃣",
  "2️⃣", "3️⃣",
  "4️⃣", "5️⃣",
  "6️⃣", "7️⃣",
  "8️⃣", "9️⃣",
  "🔟", "🔢",
  "▶️", "⏸", "⏯", "⏹", "⏺",
  "⏭", "⏮", "⏩", "⏪", "🔀",
  "🔁", "🔂", "◀️", "🔼", "🔽",
  "⏫", "⏬", "➡️", "⬅️",
  "⬆️",
  "⬇️", "↗️", "↘️", "↙️",
  "↖️",
  "↕️", "↔️", "🔄", "↪️",
  "↩️",
  "⤴️", "⤵️", "#️⃣",
  "*️⃣", "ℹ️",

  "🔤", "🔡", "🔠", "🔣", "🎵",
  "🎶", "〰️", "➰", "✔️", "🔃",
  "➕", "➖", "➗", "✖️", "💲",
  "💱", "©️", "®️", "™️",
  "🔚",
  "🔙", "🔛", "🔝", "🔜", "☑️",
  "🔘", "⚪️", "⚫️", "🔴", "🔵",
  "🔸", "🔹", "🔶", "🔷", "🔺",
  "▪️", "▫️", "⬛️",
  "⬜️", "🔻",
  "◼️", "◻️", "◾️", "◽️",
  "🔲",
  "🔳", "🔈", "🔉", "🔊", "🔇",
  "📣", "📢", "🔔", "🔕", "🃏",
  "🀄️", "♠️", "♣️", "♥️",
  "♦️",
  "🎴", "👁‍🗨", "💭", "🗯",
  "💬",
  "🕐", "🕑", "🕒", "🕓", "🕔",
  "🕕", "🕖", "🕗", "🕘", "🕙",
  "🕚", "🕛", "🕜", "🕝", "🕞",
  "🕟", "🕠", "🕡", "🕢", "🕣",
  "🕤", "🕥", "🕦", "🕧",

#   "🇦🇫", "🇦🇽", "🇦🇱",
#   "🇩🇿", "🇦🇸",
#   "🇦🇩", "🇦🇴", "🇦🇮",
#   "🇦🇶", "🇦🇬",
#   "🇦🇷", "🇦🇲", "🇦🇼",
#   "🇦🇺", "🇦🇹",
#   "🇦🇿", "🇧🇸", "🇧🇭",
#   "🇧🇩", "🇧🇧",
#   "🇧🇾", "🇧🇪", "🇧🇿",
#   "🇧🇯", "🇧🇲",
#   "🇧🇹", "🇧🇴", "🇧🇶",
#   "🇧🇦", "🇧🇼",
#   "🇧🇷", "🇮🇴", "🇻🇬",
#   "🇧🇳", "🇧🇬",
#   "🇧🇫", "🇧🇮", "🇨🇻",
#   "🇰🇭", "🇨🇲",
#   "🇨🇦", "🇮🇨", "🇰🇾",
#   "🇨🇫", "🇹🇩",
#   "🇨🇱", "🇨🇳", "🇨🇽",
#   "🇨🇨", "🇨🇴",
#   "🇰🇲", "🇨🇬", "🇨🇩",
#   "🇨🇰", "🇨🇷",
#   "🇭🇷", "🇨🇺", "🇨🇼",
#   "🇨🇾", "🇨🇿",
#   "🇩🇰", "🇩🇯", "🇩🇲",
#   "🇩🇴", "🇪🇨",
#   "🇪🇬", "🇸🇻", "🇬🇶",
#   "🇪🇷", "🇪🇪",
#   "🇪🇹", "🇪🇺", "🇫🇰",
#   "🇫🇴", "🇫🇯",
#   "🇫🇮", "🇫🇷", "🇬🇫",
#   "🇵🇫", "🇹🇫",
#   "🇬🇦", "🇬🇲", "🇬🇪",
#   "🇩🇪", "🇬🇭",
#   "🇬🇮", "🇬🇷", "🇬🇱",
#   "🇬🇩", "🇬🇵",
#   "🇬🇺", "🇬🇹", "🇬🇬",
#   "🇬🇳", "🇬🇼",
#   "🇬🇾", "🇭🇹", "🇭🇳",
#   "🇭🇰", "🇭🇺",
#   "🇮🇸", "🇮🇳", "🇮🇩",
#   "🇮🇷", "🇮🇶",
#   "🇮🇪", "🇮🇲", "🇮🇱",
#   "🇮🇹", "🇨🇮",
#   "🇯🇲", "🇯🇵", "🇯🇪",
#   "🇯🇴", "🇰🇿",
#   "🇰🇪", "🇰🇮", "🇽🇰",
#   "🇰🇼", "🇰🇬",
#   "🇱🇦", "🇱🇻", "🇱🇧",
#   "🇱🇸", "🇱🇷",
#   "🇱🇾", "🇱🇮", "🇱🇹",
#   "🇱🇺", "🇲🇴",
#   "🇲🇰", "🇲🇬", "🇲🇼",
#   "🇲🇾", "🇲🇻",
#   "🇲🇱", "🇲🇹", "🇲🇭",
#   "🇲🇶", "🇲🇷",
#   "🇲🇺", "🇾🇹", "🇲🇽",
#   "🇫🇲", "🇲🇩",
#   "🇲🇨", "🇲🇳", "🇲🇪",
#   "🇲🇸", "🇲🇦",
#   "🇲🇿", "🇲🇲", "🇳🇦",
#   "🇳🇷", "🇳🇵",
#   "🇳🇱", "🇳🇨", "🇳🇿",
#   "🇳🇮", "🇳🇪",
#   "🇳🇬", "🇳🇺", "🇳🇫",
#   "🇲🇵", "🇰🇵",
#   "🇳🇴", "🇴🇲", "🇵🇰",
#   "🇵🇼", "🇵🇸",
#   "🇵🇦", "🇵🇬", "🇵🇾",
#   "🇵🇪", "🇵🇭",
#   "🇵🇳", "🇵🇱", "🇵🇹",
#   "🇵🇷", "🇶🇦",
#   "🇷🇪", "🇷🇴", "🇷🇺",
#   "🇷🇼", "🇧🇱",
#   "🇸🇭", "🇰🇳", "🇱🇨",
#   "🇵🇲", "🇻🇨",
#   "🇼🇸", "🇸🇲", "🇸🇹",
#   "🇸🇦", "🇸🇳",
#   "🇷🇸", "🇸🇨", "🇸🇱",
#   "🇸🇬", "🇸🇽",
#   "🇸🇰", "🇸🇮", "🇸🇧",
#   "🇸🇴", "🇿🇦",
#   "🇬🇸", "🇰🇷", "🇸🇸",
#   "🇪🇸", "🇱🇰",
#   "🇸🇩", "🇸🇷", "🇸🇿",
#   "🇸🇪", "🇨🇭",
#   "🇸🇾", "🇹🇼", "🇹🇯",
#   "🇹🇿", "🇹🇭",
#   "🇹🇱", "🇹🇬", "🇹🇰",
#   "🇹🇴", "🇹🇹",
#   "🇹🇳", "🇹🇷", "🇹🇲",
#   "🇹🇨", "🇹🇻",
#   "🇺🇬", "🇺🇦", "🇦🇪",
#   "🇬🇧", "🇺🇸",
#   "🇻🇮", "🇺🇾", "🇺🇿",
#   "🇻🇺", "🇻🇦",
#   "🇻🇪", "🇻🇳", "🇼🇫",
#   "🇪🇭", "🇾🇪",
#   "🇿🇲", "🇿🇼"
]

GOING_RAW = [
    "Keshmira Vijayan",
    "Julian Garratt",
    "Sarah Lovelady",
    "Jack Yin",
    "Hirun Bandara",
    "Claire Yu",
    "Jono Kwok",
    "Caroline Ng",
    "Peter Derias",
    "Kai Mashimo",
    "Jack Vanderlaan",
    "Celine Goh",
    "Connor Pilger",
    "Amber Setchell",
    "Shrey Somaiya",
    "Capri Maher",
    "Margarita Psaras",
    "Hayes Choy",
    "Emily Tan",
    "Stanley Chen",
    "Christina Shi",
    "Angenie Bai",
    "Emily Kasovska",
    "Varisara Laosuksri",
    "Harry Partridge",
    "Chloe Cheong",
    "Christian Huang",
    "Amanda Lao",
    "Lelland Hui",
    "Emily Cong",
    "Roary Xu",
    "William Dahl",
    "Tyler Dent",
    "Lochie Bown",
    "Harry Braithwaite",
    "Alana Zhang",
    "Pravin Chanmugam",
    "Micah Wang",
    "Kenuka Wijayatunga"
]

GOING = []
for name in GOING_RAW:
    GOING.append(name.upper())

WARM_AND_FUZZIES = {}

# INDEXES
TIMESTAMP = 0
FIRST_NAME = 1
LAST_NAME = 2
MESSAGE = 3

workbook = load_workbook(filename="Warm and Fuzzies (Responses).xlsx")
sheet = workbook.active

# Go through each row until hit an empty row
for row in sheet.iter_rows(min_row=2, max_col=4, max_row=sheet.max_row):
    name = f"{row[FIRST_NAME].value.strip()} {row[LAST_NAME].value.strip()}".upper()

    if name not in WARM_AND_FUZZIES:
        WARM_AND_FUZZIES[name] = []

    WARM_AND_FUZZIES[name].append(row[MESSAGE].value)

# print(json.dumps(WARM_AND_FUZZIES, indent=2))

# print(len(WARM_AND_FUZZIES), "recipients...")

# Dump each person into a LaTeX file
count = 0
for name in WARM_AND_FUZZIES:
    doc = Document()

    doc.styles["Heading 1"].font.name = "CMU Serif"
    doc.styles["Heading 1"].font.size = Pt(24)
    doc.styles["Heading 1"].font.color.rgb = RGBColor(0, 0, 0)

    doc.styles["Normal"].font.name = "CMU Serif"
    doc.styles["Normal"].paragraph_format.space_before = Pt(12)

    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run("")
    run.add_picture("logo.png", width=Inches(0.5))

    doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = 1
    run = t.add_run(name)
    run.font.size = Pt(24)
    run.bold = True

    h = doc.add_paragraph()
    h.alignment = 1
    h.add_run("Co-op Soc 2020").italic = True

    for message in WARM_AND_FUZZIES[name]:
        divider = doc.add_paragraph(random.choice(EMOJIS))
        divider.alignment = 1
        doc.add_paragraph(message.strip().strip("\n"))

    if name in GOING:
        count += 1
        file_name = "going_" + name.replace(" ", "_") + ".docx"
    else:
        file_name = name.replace(" ", "_") + ".docx"

    doc.save(file_name)

print(count)